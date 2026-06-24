"""
docx_generator.py
COA Word doc — mirrors the PDF layout:
  2-col header (logo | contact+address) → product info grid → COA title →
  bordered test-results table (4 cols, category bold rows) → signature block
"""
from __future__ import annotations
import os
from datetime import date
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from core.coa_model import COADocument, PassFail, TestCategory, build_display_rows, fit_scale

# Best-effort single-page fit scale for the document currently being built —
# set once at the top of generate_docx() and read by _cell_para/_add_para.
_scale = 1.0

BLACK = RGBColor(0x00, 0x00, 0x00)
DGRAY = RGBColor(0x55, 0x55, 0x55)
LGRAY = RGBColor(0xF0, 0xF0, 0xF0)
GREEN = RGBColor(0x1A, 0x7A, 0x4A)
RED   = RGBColor(0xC0, 0x39, 0x2B)
NAVY  = RGBColor(0x0D, 0x2B, 0x55)

_CATEGORY_LABELS = {
    TestCategory.PHYSICAL.value:        "Physical Characteristics",
    TestCategory.CHEMICAL.value:        "Chemical Properties",
    TestCategory.MICROBIOLOGICAL.value: "Microbiological Data",
    TestCategory.OTHER.value:           "Other",
}


def _hex(rgb: RGBColor) -> str:
    return f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}"


def _set_cell_bg(cell, rgb: RGBColor):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), _hex(rgb))
    tcPr.append(shd)


def _cell_para(cell, text: str, bold=False, italic=False,
               color: RGBColor | None = None, font_size: float = 9,
               align=WD_ALIGN_PARAGRAPH.LEFT) -> None:
    cell.paragraphs[0].clear()
    para = cell.paragraphs[0]
    para.alignment = align
    if text:
        run = para.add_run(text)
        run.bold = bold
        run.italic = italic
        run.font.size = Pt(font_size * _scale)
        if color:
            run.font.color.rgb = color


def _add_bottom_border(cell, color_hex="000000", size=6):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single")
    bot.set(qn("w:sz"), str(size))
    bot.set(qn("w:color"), color_hex)
    tcBorders.append(bot)
    tcPr.append(tcBorders)


def _add_para(doc: Document, text: str, bold=False, italic=False,
              font_size: float = 9, color: RGBColor | None = None,
              align=WD_ALIGN_PARAGRAPH.LEFT, space_after: float = 0) -> None:
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    if text:
        r = p.add_run(text)
        r.bold = bold
        r.italic = italic
        r.font.size = Pt(font_size * _scale)
        if color:
            r.font.color.rgb = color


def _add_hr(doc: Document):
    """Add a horizontal rule paragraph."""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "888888")
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)


def generate_docx(coa: COADocument, output_path: str) -> None:
    global _scale
    _scale = fit_scale(len(coa.test_results))

    doc = Document()

    for section in doc.sections:
        section.top_margin    = Cm(1.8)
        section.bottom_margin = Cm(1.8)
        section.left_margin   = Cm(1.8)
        section.right_margin  = Cm(1.8)

    # Shared total table width so every table's right edge lines up.
    sec = doc.sections[0]
    usable_w = sec.page_width - sec.left_margin - sec.right_margin

    # ── 1. COMPANY HEADER ─────────────────────────────────────────────────────
    if coa.receiving_company_header_path and os.path.exists(coa.receiving_company_header_path):
        try:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run().add_picture(coa.receiving_company_header_path, width=Inches(6.3))
        except Exception:
            pass
    else:
        ht = doc.add_table(rows=1, cols=2)
        lc = ht.cell(0, 0)
        rc = ht.cell(0, 1)

        # Left: logo or company name
        if coa.receiving_company_logo_path and os.path.exists(coa.receiving_company_logo_path):
            try:
                p = lc.paragraphs[0]
                p.add_run().add_picture(coa.receiving_company_logo_path, width=Cm(4))
            except Exception:
                _cell_para(lc, coa.receiving_company_name or "", bold=True, font_size=12, color=NAVY)
        else:
            _cell_para(lc, coa.receiving_company_name or "", bold=True, font_size=12, color=NAVY)

        # Right: company contact + address
        right_lines = []
        if coa.receiving_company_phone:
            right_lines.append(f"Tel: {coa.receiving_company_phone}")
        if coa.receiving_company_website:
            right_lines.append(f"Web: {coa.receiving_company_website}")
        addr_joined = ", ".join(p.strip() for p in (coa.receiving_company_address or "").split("\n") if p.strip())
        if addr_joined:
            right_lines.append(f"Add: {addr_joined}")

        if right_lines:
            _cell_para(rc, right_lines[0], font_size=8, align=WD_ALIGN_PARAGRAPH.RIGHT)
            for line in right_lines[1:]:
                p2 = rc.add_paragraph(line)
                p2.runs[0].font.size = Pt(8)
                p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        for cell, w in zip([lc, rc], [int(usable_w * 0.4), int(usable_w * 0.6)]):
            cell.width = w

    # ── 4. COA TITLE ─────────────────────────────────────────────────────────
    _add_para(doc, "CERTIFICATE OF ANALYSIS", bold=True, font_size=14,
              align=WD_ALIGN_PARAGRAPH.CENTER)

    # ── 2. (removed — no supplier/subject company block) ─────────────────────

    # ── 3. PRODUCT INFO GRID ──────────────────────────────────────────────────
    grid_pairs = [
        (("Product name", coa.product_name), ("Batch no.", coa.lot_number)),
        (("Botanical name", coa.botanical_name), ("Production date", coa.manufacturing_date)),
        (("Plant part", coa.plant_part), ("Analysis date", coa.date_of_analysis)),
        (("Country of origin", coa.manufacturer_country), ("Re-test date", coa.retest_date)),
    ]
    grid_rows = [
        (l1, v1, l2, v2) for (l1, v1), (l2, v2) in grid_pairs
        if (v1 and str(v1).strip()) or (v2 and str(v2).strip())
    ]

    if grid_rows:
        it = doc.add_table(rows=len(grid_rows), cols=4)
        col_widths = [int(usable_w * 0.16), int(usable_w * 0.34), int(usable_w * 0.16), int(usable_w * 0.34)]
        for i, (l1, v1, l2, v2) in enumerate(grid_rows):
            cells = it.row_cells(i)
            _cell_para(cells[0], l1, bold=True, font_size=9)
            _cell_para(cells[1], str(v1 or ""), font_size=9)
            _cell_para(cells[2], l2, bold=True, font_size=9)
            _cell_para(cells[3], str(v2 or ""), font_size=9)
            for cell, w in zip(cells, col_widths):
                cell.width = w

    doc.add_paragraph()

    # ── 5. TEST RESULTS TABLE ─────────────────────────────────────────────────
    if coa.test_results:
        display_rows = build_display_rows(coa.test_results, _CATEGORY_LABELS)

        total = 1 + len(display_rows)
        t = doc.add_table(rows=total, cols=4)
        t.style = "Table Grid"
        col_widths = [int(usable_w * 0.28), int(usable_w * 0.24), int(usable_w * 0.24), int(usable_w * 0.24)]
        hdrs = ["Analysis Item", "Specification", "Result", "Analysis Test Method"]

        # Header row
        hdr_cells = t.row_cells(0)
        for j, (cell, h, w) in enumerate(zip(hdr_cells, hdrs, col_widths)):
            _cell_para(cell, h, bold=True, font_size=9)
            _add_bottom_border(cell, "000000", 12)
            cell.width = w

        row_idx = 1
        for kind, payload in display_rows:
            if kind == "band":
                cat_cells = t.row_cells(row_idx)
                cat_cells[0].merge(cat_cells[3])
                _set_cell_bg(cat_cells[0], LGRAY)
                _cell_para(cat_cells[0], payload, bold=True, font_size=9)
            else:
                tr = payload
                cells = t.row_cells(row_idx)
                result_text = f"{tr.result} {tr.unit}".strip() if tr.unit else (tr.result or "")
                result_color = (
                    GREEN if tr.pass_fail == PassFail.PASS.value
                    else RED if tr.pass_fail == PassFail.FAIL.value
                    else None
                )
                _cell_para(cells[0], tr.test_name or "", font_size=8.5)
                _cell_para(cells[1], tr.specification or "", font_size=8.5)
                _cell_para(cells[2], result_text, font_size=8.5, color=result_color)
                _cell_para(cells[3], tr.method or "", font_size=8.5)
                for cell, w in zip(cells, col_widths):
                    cell.width = w
            row_idx += 1

    else:
        doc.add_paragraph("No test results recorded.")

    doc.add_paragraph()

    # ── 7. SIGNATURE BLOCK ────────────────────────────────────────────────────
    sig_date = coa.signature_date or date.today().strftime("%Y-%m-%d")
    _add_para(doc, "Approved by QC Supervisor", font_size=9, color=DGRAY, space_after=2)
    _add_para(doc, sig_date, font_size=9)

    # ── 8. FOOTER ─────────────────────────────────────────────────────────────
    # A true Word page footer (doc.sections[0].footer) instead of a body
    # paragraph — so it repeats at the bottom of every page automatically,
    # rather than landing wherever the body content happens to end.
    footer_text = (
        f"Certificate of Analysis — {coa.product_name} — Lot: {coa.lot_number} — "
        f"Issued: {date.today().strftime('%m/%d/%Y')} — "
        "Results relate only to the sample as received. "
        "This document shall not be reproduced except in full without written approval."
    )
    footer = doc.sections[0].footer
    footer.is_linked_to_previous = False
    footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    footer_para.text = ""
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pPr = footer_para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    top = OxmlElement("w:top")
    top.set(qn("w:val"), "single")
    top.set(qn("w:sz"), "4")
    top.set(qn("w:space"), "4")
    top.set(qn("w:color"), "DDDDDD")
    pBdr.append(top)
    pPr.append(pBdr)
    run = footer_para.add_run(footer_text)
    run.font.size = Pt(7.5 * _scale)
    run.font.color.rgb = DGRAY

    doc.save(output_path)
