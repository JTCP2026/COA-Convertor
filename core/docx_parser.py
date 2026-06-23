"""
docx_parser.py
Extracts text and tables from .docx files using python-docx.
"""
from __future__ import annotations
import logging
from .pdf_parser import ParsedDocument

log = logging.getLogger(__name__)


def parse_docx(path: str, progress_callback=None) -> ParsedDocument:
    from docx import Document

    result = ParsedDocument(source_type="docx")
    text_parts: list[str] = []
    tables: list[list[list[str]]] = []

    try:
        doc = Document(path)

        total = len(doc.paragraphs) + len(doc.tables)
        done = 0

        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
            done += 1
            if progress_callback and total:
                progress_callback(int((done / total) * 90))

        for table in doc.tables:
            rows: list[list[str]] = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append(cells)
            if rows:
                tables.append(rows)
            done += 1
            if progress_callback and total:
                progress_callback(int((done / total) * 90))

        result.text = "\n".join(text_parts)
        result.tables = tables
        result.page_count = 1

    except Exception as e:
        log.error("docx parse error: %s", e)

    if progress_callback:
        progress_callback(95)

    return result
